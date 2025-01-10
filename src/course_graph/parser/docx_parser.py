# -*- coding: utf-8 -*-
# Create Date: 2024/07/11
# Author: wangtao <wangtao.cpu@gmail.com>
# File Name: course_graph/parser/docx_parser.py
# Description: 定义docx文档解析器

from .types import BookMark, PageIndex
from .parser import Parser
from .types import Content, ContentType
import docx
import shortuuid
from xml.dom.minidom import parseString


class DOCXParser(Parser):

    def __init__(self, docx_path: str) -> None:
        """ 解析docx格式文档, 需要带有大纲级别以判断层级

        Args:
            docx_path(str): docx文档路径
        """
        super().__init__(docx_path)
        self.__docx = docx.Document(docx_path)

    def __enter__(self) -> 'DOCXParser':
        return self

    def close(self) -> None:
        pass

    def get_bookmarks(self) -> list[BookMark]:
        stack: list[BookMark] = []
        bookmarks: list[BookMark] = []
        for phar in self.__docx.paragraphs:
            parser = parseString(phar._p.xml)
            if len(element := parser.getElementsByTagName('w:outlineLvl')) > 0:
                level = int(element[0].getAttribute('w:val'))
                title = phar.text
                bookmarks.append(
                    BookMark(
                        id='1:' + str(shortuuid.uuid()) + f':{level}',
                        title=title,
                        page_start=PageIndex(index=0,
                                             anchor=(0, 0)),  # 无需设置起始页码和结束页码
                        page_end=PageIndex(index=0, anchor=(0, 0)),
                        level=level,
                        subs=[],
                        resource=[]))
        # 先获取全部的书签再合并
        for bookmark in reversed(bookmarks):
            level = bookmark.level

            while stack and stack[-1].level > level:
                bookmark.subs.append(stack.pop())

            stack.append(bookmark)

        stack.reverse()
        return stack

    def get_contents(self, bookmark: BookMark) -> list[Content]:
        contents: list[Content] = []
        start = False
        for phar in self.__docx.paragraphs:
            if len(text := phar.text) == 0:
                continue
            parser = parseString(phar._p.xml)
            if len(element := parser.getElementsByTagName('w:outlineLvl')) > 0:
                level = int(element[0].getAttribute('w:val'))
            else:
                level = -1
            if level <= bookmark.level and level != -1 and start:
                break
            if level == bookmark.level and text == bookmark.title:
                start = True
            if start:
                contents.append(
                    Content(type=ContentType.Title
                            if level != -1 else ContentType.Text,
                            content=text))
        return contents
